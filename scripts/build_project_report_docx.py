from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RebootRoute_Project_Report.docx"
FONT = "Apple SD Gothic Neo"
ACCENT = RGBColor(15, 118, 110)
INK = RGBColor(23, 33, 29)
MUTED = RGBColor(89, 101, 96)
HEADER_FILL = "E8EEF5"
SOFT_FILL = "F6F8F5"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.25) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p)
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, color=INK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=ACCENT if level <= 2 else INK)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        set_paragraph_spacing(p, after=4)
        run = p.add_run(item)
        set_run_font(run, color=INK)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        set_paragraph_spacing(p, after=4)
        run = p.add_run(item)
        set_run_font(run, color=INK)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=9.2, bold=bold, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    repeat_table_header(table.rows[0])
    keep_row_together(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_width(table.rows[0].cells[idx], widths[idx])
        set_cell_text(table.rows[0].cells[idx], header, bold=True, fill=HEADER_FILL)
    for row in rows:
        table_row = table.add_row()
        keep_row_together(table_row)
        cells = table_row.cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_text(cells[idx], value)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def add_callout(doc: Document, title: str, body: str, fill: str = SOFT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 9360)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.15)
    p.text = ""
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.15)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.8, color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def read_metadata() -> dict:
    path = ROOT / "models" / "latest" / "metadata.json"
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def build_doc() -> None:
    metadata = read_metadata()
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=3)
    run = title.add_run("RebootRoute 프로젝트 발표·구현 명세서")
    set_run_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    run = subtitle.add_run("인천 청년정책·문화활동 RAG + 부담도 기반 미션 추천 MVP")
    set_run_font(run, size=13, color=ACCENT)

    add_callout(
        doc,
        "문서 목적",
        "이 문서는 발표자가 RebootRoute의 구현 범위, 실제 데이터 구성, 대시보드 사용 순서, 코드 구조, MLOps 산출물, 한계와 운영 전 교체 항목을 정확히 설명할 수 있도록 정리한 프로젝트 명세서입니다.",
    )
    add_table(
        doc,
        ["항목", "내용"],
        [
            ["프로젝트명", "RebootRoute"],
            ["MVP 범위", "인천 청년정책/문화활동 RAG + 부담도 기반 단계형 미션 추천"],
            ["생성일", date.today().isoformat()],
            ["대상", "인천 거주 19-39세 고립·은둔 또는 고립 위험 청년, 가족·지인, 지역 청년지원기관/문화기관 운영자"],
            ["중요 경고", "진단, 치료, 상담 챗봇, 위험도 판정, 단순 취업 추천 앱이 아님"],
        ],
        [2100, 7260],
    )

    doc.add_page_break()

    add_heading(doc, "1. 프로젝트 정의", 1)
    add_paragraph(
        doc,
        "RebootRoute는 인천의 실제 청년정책·문화공간·문화행사 정보를 공식 출처 기반으로 보여주고, 사용자가 오늘 바로 끝낼 수 있는 작은 다음 행동을 제안하는 MVP입니다. 현재 실제 사용자 데이터를 확보할 수 없으므로, 사용자 화면은 개인 상태 입력을 요구하지 않고 공식 자원 확인과 조건 축소 흐름을 중심으로 설계했습니다.",
    )
    add_bullets(
        doc,
        [
            "현재 핵심 축은 공식 출처 기반 resource seed, local TF-IDF RAG, 부담도 기반 다음 행동 제안입니다.",
            "서비스 문구는 진단/치료/상담이 아니라 정보 탐색과 낮은 부담 행동 제안으로 제한합니다.",
            "자해·타해·즉시 위험 표현이 들어오면 일반 추천을 중단하고 안전 자원 연결로 분기합니다.",
            "Synthetic profile과 synthetic label은 모델 학습, API 호환성, 운영자 점검, MLOps 시연 용도로만 유지합니다.",
            "초기 검증은 closed-loop A/B가 아니라 open-loop evaluation, human rubric evaluation, batch retraining이 적절합니다.",
        ],
    )

    add_heading(doc, "2. 사용자 화면 사용 순서", 1)
    add_paragraph(
        doc,
        "대시보드의 첫 탭은 실제 사용자가 무엇을 먼저 보고 무엇을 얻는지 명확히 알 수 있도록 네 단계로 재구성했습니다. 조건을 변경하면 추천 루트, 미션, 자원 후보, 지도 위치가 즉시 갱신되며, 별도 업데이트 버튼이나 추천 루트 버튼을 누르지 않습니다.",
    )
    add_table(
        doc,
        ["순서", "화면에서 하는 일", "사용자가 얻는 것"],
        [
            ["1", "내 조건 입력", "위치, 외출 가능 시간, 비용, 대면 부담, 관심 분야"],
            ["2", "추천 루트 확인", "현재 부담도에 맞는 단계와 오늘 할 미션"],
            ["3", "지도 확인", "내 위치와 활동 장소의 상대 위치, 대략 직선거리"],
            ["4", "결과 기록", "미션 시작/완료/too-hard, 프로그램 참여, 지원 신청/결과, 미니 프로젝트 제출"],
        ],
        [900, 3200, 5260],
    )
    add_callout(
        doc,
        "사용자 화면에서 숨기는 정보",
        "resource_id, mission_id, ranking score, RAG score, reboot point, model metric, raw payload는 사용자 카드에 표시하지 않습니다. 이 정보는 검증 탭과 API 응답에서만 내부 검증용으로 확인합니다.",
        fill="FFF3D8",
    )

    doc.add_page_break()
    add_heading(doc, "3. 전체 아키텍처", 1)
    add_table(
        doc,
        ["순서", "구성 요소", "구현 위치", "역할"],
        [
            ["1", "Sample data + official resource seed", "src/rebootroute/data/mock_data.py, data/raw/*.csv", "Synthetic profile/progress와 공식 출처 기반 인천 자원 seed 생성"],
            ["2", "Data Validation", "src/rebootroute/data/validation.py", "필수 column, 타입, 범위 검증. pandera 미설치 시 fallback"],
            ["3", "Feature Build", "src/rebootroute/features/build_features.py", "부담도, readiness, progress, interest feature와 synthetic label 생성"],
            ["4", "RAG Resource Index", "src/rebootroute/rag/retriever.py", "TF-IDF 기반 공식 자원 검색"],
            ["5", "Safety Guardrail", "src/rebootroute/recommender/safety_guardrails.py", "위험 표현 감지 시 안전 자원 분기"],
            ["6", "Stage + Mission 추천", "src/rebootroute/recommender/*.py", "rule stage, mission ranking, resource ranking, mini project 생성"],
            ["7", "API / Dashboard", "src/rebootroute/api/main.py, src/rebootroute/dashboard/app.py", "FastAPI와 Streamlit 오늘 루트/자원·지도/기록/검증 화면"],
            ["8", "Operational Logs", "src/rebootroute/database.py", "progress_logs, feedback_events, outcome_events, user_state 저장"],
            ["9", "Evaluation / Retraining", "evaluation/*.csv, scripts/build_human_eval_sheet.py, scripts/run_pipeline.py", "human eval sheet와 batch 학습 루프"],
        ],
        [760, 1720, 3320, 3560],
    )
    doc.add_page_break()

    add_heading(doc, "4. 코드 구조", 1)
    add_table(
        doc,
        ["경로", "내용"],
        [
            ["configs/config.yaml", "프로젝트명, 경로, 추천 top_n, safety resource path 설정"],
            ["configs/safety_resources.example.yaml", "위험 표현 분기 시 보여줄 도움 연결 자원"],
            ["src/rebootroute/schemas.py", "Pydantic schema: UserProfile, Mission, Resource, ProgressLog, FeedbackEvent, OutcomeEvent, RAGSearchRequest"],
            ["src/rebootroute/database.py", "SQLite 연결, DB 초기화, progress/feedback/outcome 저장, reboot point 조회"],
            ["src/rebootroute/api/main.py", "FastAPI endpoint 정의"],
            ["src/rebootroute/dashboard/app.py", "Streamlit 대시보드. 오늘 루트, 자원·지도, 기록, 검증 탭"],
            ["src/rebootroute/data/*.py", "Synthetic sample 생성, 공식 resource seed 생성, validation"],
            ["src/rebootroute/features/build_features.py", "feature table과 synthetic label 생성"],
            ["src/rebootroute/modeling/*.py", "모델 학습, 평가, registry, prediction, explanation"],
            ["src/rebootroute/rag/retriever.py", "TF-IDF local retrieval"],
            ["src/rebootroute/recommender/*.py", "stage rule, mission/resource ranking, route builder, safety, mini project"],
            ["scripts/*.py", "pipeline, sample data, DB init, model training, human eval sheet 생성"],
            ["tests/*.py", "validation, stage rules, recommender, RAG, feedback, safety, API smoke test"],
        ],
        [3500, 5860],
    )
    doc.add_page_break()

    add_heading(doc, "5. 데이터와 저장소 산출물", 1)
    add_table(
        doc,
        ["파일/테이블", "내용", "주의"],
        [
            ["data/raw/sample_profiles.csv", "180개 synthetic profile", "실제 사용자가 아님"],
            ["data/raw/sample_missions.csv", "Stage 0-7 미션 42개", "발표용 seed data"],
            ["data/raw/sample_resources.csv", "공식 출처 기반 인천 자원 seed", "인천청년포털, 인천문화재단, 인천아트플랫폼, 트라이보울 등"],
            ["data/raw/sample_progress.csv", "synthetic 진행 로그", "실제 완료 이력이 아님"],
            ["data/raw/sample_outcomes.csv", "outcome import template", "실제 outcome row는 사용자 또는 기관 관측 후 입력"],
            ["data/features/training_features.csv", "학습 feature + synthetic labels", "production 전 교체 필요"],
            ["data/rebootroute.db", "SQLite DB", "feedback_events, progress_logs, outcome_events, user_state"],
            ["models/latest/*.joblib", "stage / mission success model", "MLOps 시연용"],
            ["models/latest/metadata.json", "모델명, metric, data_version, warning", "dashboard MLOps 탭에서 사용"],
            ["reports/data_card.md", "데이터 설명과 synthetic label 경고", "발표/검토 산출물"],
            ["reports/model_card.md", "모델 후보, 선택 모델, metric, 한계", "발표/검토 산출물"],
            ["reports/human_eval_review_sheet.csv", "40개 human eval case 결과 sheet", "사람이 빈 score column을 채점"],
        ],
        [2650, 4210, 2500],
    )

    add_heading(doc, "6. 추천 구현 상세", 1)
    add_heading(doc, "6.1 사용자 화면 정렬", 2)
    add_paragraph(
        doc,
        "현재 사용자 화면은 데모 세션 profile과 공식 자원 속성을 함께 사용합니다. 비용, 부담도, 온라인 확인 가능 여부, 내 위치와의 직선거리, 예상 확인 시간을 기준으로 자원을 정렬하고, 가장 부담이 낮은 후보를 기준으로 오늘의 다음 행동을 만듭니다.",
    )
    add_bullets(
        doc,
        [
            "정렬 기준: burden_level 낮음, cost_type 낮음, online_available 우선, estimated_duration_minutes 짧음, 이름 순",
            "다음 행동: 공식 페이지 열기, 대상·비용·운영시간 확인, 문의처 메모, 조건 1줄 정리",
            "사용자 카드에는 내부 score와 ID를 표시하지 않습니다.",
        ],
    )
    add_heading(doc, "6.2 Rule 기반 stage 분류", 2)
    add_paragraph(
        doc,
        "stage_rules.py는 API와 운영자 점검용으로 남아 있습니다. 외출 부담, 대면 부담, 에너지, 생활 리듬, 취업 부담, 지원자 존재 여부, 과거 진행 로그를 사용해 Stage 0-7을 정합니다. 추천의 1차 기준은 rule이며 ML stage model은 보조 stage와 contributing factor 설명용입니다.",
    )
    add_heading(doc, "6.3 Mission ranking", 2)
    add_bullets(
        doc,
        [
            "stage_match: 추천 stage와 미션 stage가 가까울수록 높은 점수",
            "burden_fit: 사용자의 에너지/생활 리듬과 외출/대면 부담을 고려",
            "interest_fit: 관심 태그와 career/recovery tag overlap",
            "resource_available: 연결 가능한 자원이 있을 때 보정",
            "progress_continuity: 너무 어려움 기록이 있으면 보수적으로 조정",
            "novelty: 완료/건너뜀/너무 어려움 이력을 반영",
            "cost_fit: 현재 mission에는 직접 비용이 없으므로 budget readiness signal로 사용",
        ],
    )
    add_heading(doc, "6.4 Resource ranking", 2)
    add_bullets(
        doc,
        [
            "district_match: 사용자 구/군과 자원 구/군 일치",
            "burden_fit: outdoor/social 부담과 사용자 capacity 비교",
            "interest_fit: 관심 태그와 resource tag overlap",
            "contact_mode_fit: 온라인/낮은 대면/소규모/대면 가능 선호 반영",
            "cost_fit: 무료/저비용/유료와 budget limit 비교",
            "career_fit: 직무탐색 단계 이상에서 career tag를 추가 반영",
            "이 점수는 운영자 점검용이며 사용자 화면에는 표시하지 않습니다.",
        ],
    )

    add_heading(doc, "7. RAG 구현 상세", 1)
    add_paragraph(
        doc,
        "retriever.py는 sample_resources.csv를 읽어 resource name, description, district, type, cost, career/recovery tag를 하나의 text로 묶고, scikit-learn TfidfVectorizer로 query와 resource 간 cosine similarity를 계산합니다. analyzer는 char_wb, ngram_range는 2-4입니다. sklearn이 없으면 query token overlap fallback을 사용합니다.",
    )
    add_bullets(
        doc,
        [
            "기본 검색 대상은 youth_program, culture_event, culture_facility, support_program, contest, mini_project입니다.",
            "district가 주어지면 같은 구/군 자원을 앞쪽에 재정렬합니다.",
            "max_burden_level이 있으면 해당 부담도 이하만 남깁니다.",
            "반환 answer에는 진단/치료가 아니라 정보 탐색 참고 자료라는 안전 문구를 포함합니다.",
            "사용자 화면에서는 자원명, 설명, 구/군, 비용, 부담도, 주소, 문의처, 공식 출처 링크만 보여주고 score/id/source key는 숨깁니다.",
        ],
    )

    add_heading(doc, "8. Safety guardrail", 1)
    add_paragraph(
        doc,
        "safety_guardrails.py는 자해, 타해, 즉시 위험 표현을 regex로 감지합니다. 감지되면 route_builder.analyze_profile은 recommended_stage 0, safety_flag True, risk_type urgent_support_needed를 반환하고 next_3_missions와 recommended_resources를 빈 배열로 둡니다.",
    )
    add_table(
        doc,
        ["항목", "구현"],
        [
            ["감지 예", "죽고 싶, 자살, 극단적 선택, 사라지고 싶, 해치고 싶, kill myself, hurt someone 등"],
            ["분기 결과", "일반 미션 추천 중단, 안전 확인 메시지, safety resources 반환"],
            ["설정 파일", "configs/safety_resources.example.yaml"],
            ["테스트", "tests/test_safety_guardrails.py"],
        ],
        [2200, 7160],
    )

    add_heading(doc, "9. Feedback loop", 1)
    add_paragraph(
        doc,
        "현재 운영 전 데이터 루프 설계는 상태 입력 → safety guardrail → stage 분류 → RAG/mission/resource 추천 → progress/feedback/outcome log → human evaluation → batch retraining입니다. 실제 사용자 데이터가 없어 학습 label은 synthetic placeholder이지만, 데모 중 미션 시작/완료/too-hard와 프로그램 참여, 지원 신청, 지원 결과, 미니 프로젝트 제출, 운영자 검토는 SQLite에 기록됩니다.",
    )
    add_table(
        doc,
        ["테이블", "주요 필드", "역할"],
        [
            ["progress_logs", "log_id, user_id, mission_id, status, completed_at, points_awarded, payload_json", "미션 시작/완료/건너뜀/너무 어려움 진행 로그"],
            ["feedback_events", "event_id, user_id, event_type, mission_id, resource_id, rating, user_note, policy_version", "추천 반응과 운영자 review event 저장"],
            ["outcome_events", "outcome_id, user_id, outcome_type, outcome_status, mission_id, resource_id, readiness_rating, burden_after, result_note", "프로그램 참여, 지원 신청/결과, 미니 프로젝트 제출, 운영자 검토 outcome 저장"],
            ["user_state", "user_id, reboot_points, last_stage, updated_at", "사용자별 내부 상태 저장. point는 연구/운영용이며 사용자 화면에는 노출하지 않음"],
        ],
        [1700, 4700, 2960],
    )

    doc.add_page_break()
    add_heading(doc, "10. Dashboard 재설계 내용", 1)
    add_table(
        doc,
        ["탭", "보여주는 대상", "내용"],
        [
            ["오늘 루트", "실제 사용자", "내 조건 입력 → 추천 루트 → 지도 → 미션/결과 기록. ID, score, point 숨김"],
            ["자원·지도", "실제 사용자", "RAG 검색 결과를 자원 카드, 공식 출처 링크, 내 위치/활동 장소 지도 중심으로 표시"],
            ["기록", "실제 사용자/발표자", "데모 세션의 progress/outcome/feedback 개수와 최근 로그 확인"],
            ["검증", "발표자/운영자", "Rule stage, ML 보조 stage, contributing factors, IDs, score, raw payload, feedback/progress/outcome log, 운영자 review 입력, 모델 metric 확인"],
        ],
        [1700, 1900, 5760],
    )
    add_bullets(
        doc,
        [
            "첫 화면은 내 조건, 추천 루트, 지도, 결과 기록 순서로 보이도록 재구성했습니다.",
            "조건 변경 즉시 추천 루트, 미션, 자원 후보, 지도 위치가 갱신됩니다.",
            "사용자 화면의 행동은 시작, 완료, 너무 어려움, 참여/지원 결과처럼 실제 학습 loop에 필요한 기록으로 정리했습니다.",
            "사용자 화면에는 내부 ranking score, 식별자, point를 숨기고, 검증 탭에서만 표로 공개합니다.",
            "위도/경도 기반 미니맵으로 내 위치와 활동 장소를 함께 표시합니다.",
            "미션 시작/완료/too-hard, 프로그램 참여, 지원 신청/결과, 미니 프로젝트 제출, 운영자 review를 SQLite에 저장합니다.",
            "모바일 폭에서는 단계 안내와 필터/카드가 한 열로 쌓이도록 반응형 CSS를 적용했습니다.",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "11. API 구현", 1)
    add_table(
        doc,
        ["Endpoint", "Method", "역할"],
        [
            ["/health", "GET", "서비스 상태"],
            ["/metadata", "GET", "모델 metadata 반환"],
            ["/sample_profile", "GET", "샘플 profile 반환"],
            ["/analyze_intake", "POST", "intake 분석, safety, stage, top missions/resources 반환"],
            ["/recommend_route", "POST", "현재 stage부터 최대 3개 stage route 반환"],
            ["/recommend_missions", "POST", "미션 ranking 반환"],
            ["/recommend_resources", "POST", "자원 ranking 반환"],
            ["/rag/search", "POST", "정책/문화 RAG 검색"],
            ["/feedback/log", "POST", "feedback event 저장"],
            ["/progress/log", "POST", "progress log 저장 및 point 갱신"],
            ["/outcomes/log", "POST", "program participation, support application/result, mini project submission, operator review outcome 저장"],
            ["/outcomes", "GET", "사용자별 outcome log 조회"],
            ["/simulate", "POST", "field 변화 시 stage 변화 simulation"],
        ],
        [2500, 1100, 5760],
    )

    add_heading(doc, "12. MLOps와 평가", 1)
    stage_model = metadata.get("stage_model_name", "unknown")
    mission_model = metadata.get("mission_success_model_name", "unknown")
    stage_acc = metadata.get("stage_metrics", {}).get("accuracy")
    stage_f1 = metadata.get("stage_metrics", {}).get("macro_f1")
    mission_acc = metadata.get("mission_success_metrics", {}).get("accuracy")
    mission_f1 = metadata.get("mission_success_metrics", {}).get("macro_f1")
    mission_auc = metadata.get("mission_success_metrics", {}).get("roc_auc")
    add_table(
        doc,
        ["항목", "현재 값"],
        [
            ["Stage 선택 모델", str(stage_model)],
            ["Mission success 선택 모델", str(mission_model)],
            ["Data version", str(metadata.get("data_version", "unknown"))],
            ["Stage accuracy / macro F1", f"{stage_acc} / {stage_f1}"],
            ["Mission accuracy / macro F1 / ROC-AUC", f"{mission_acc} / {mission_f1} / {mission_auc}"],
            ["MLflow tracking", "data/mlflow.db SQLite backend, 실패 시 mlruns/fallback_runs.jsonl"],
            ["Human eval cases", "evaluation/human_eval_cases.csv 40개"],
            ["Rubric", "stage_score, mission_burden_score, rag_grounding_score, safety_score"],
        ],
        [3200, 6160],
    )
    add_callout(
        doc,
        "Synthetic label warning",
        "현재 학습 label은 사용자 행동과 기관 결과를 대신한 synthetic placeholder입니다. 로그 수집, 운영자 검토, 재학습 구조는 구현되어 있지만, 실제 미션 완료, too-hard 피드백, 프로그램 참여, 지원 결과처럼 관측이 필요한 outcome label은 운영 또는 기관 연계 후 수집·import해야 합니다.",
        fill="FFF3D8",
    )

    add_heading(doc, "13. 실행 및 검증 명령", 1)
    add_table(
        doc,
        ["명령", "설명"],
        [
            ["make setup", "requirements.txt 의존성 설치"],
            ["make pipeline", "synthetic sample과 공식 resource seed 생성, validation, feature build, model training, DB init, reports 생성"],
            ["make dashboard", "Streamlit dashboard 실행. 기본 http://localhost:8501, 포트 사용 중이면 터미널 Local URL 사용"],
            ["make api", "FastAPI 실행, http://localhost:8000/docs"],
            ["make test", "pytest 실행"],
            ["make eval-sheet", "reports/human_eval_review_sheet.csv 생성"],
            ["make docker-up", "docker compose up --build"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "14. 발표 데모 시나리오", 1)
    add_numbers(
        doc,
        [
            "make dashboard로 Streamlit을 실행한다.",
            "오늘 루트 탭에서 내 조건 → 추천 루트 → 지도 → 결과 기록 순서가 보이는지 확인한다.",
            "외출 부담, 대면 부담, 위치, 자원 종류, 비용, 최대 부담도를 바꿔 추천 미션과 자원 후보가 즉시 바뀌는 것을 보여준다.",
            "지도에서 내 위치와 활동 장소가 함께 표시되는 것을 보여준다.",
            "미션 시작/완료/너무 어려움 버튼을 눌러 progress log가 저장되는 것을 보여준다.",
            "활동/지원 결과 기록 폼에서 프로그램 참여 또는 지원 신청 outcome을 저장한다.",
            "자원·지도 탭에서 RAG 검색 결과, 공식 출처 링크, 지도 표시를 보여준다.",
            "기록 탭에서 progress/outcome/feedback 로그가 쌓인 것을 확인한다.",
            "검증 탭에서 mission_id, resource_id, score, raw payload, feedback/outcome log가 내부 검증용으로만 보이는 것을 확인한다.",
            "API 문서에서 safety guardrail, feedback, progress, outcomes endpoint가 구현되어 있음을 보여준다.",
        ],
    )

    add_heading(doc, "15. 운영 전 교체·보완 항목", 1)
    add_table(
        doc,
        ["항목", "현재 상태", "운영 전 필요한 작업"],
        [
            ["공공데이터", "공식 출처 기반 curated seed", "자동 동기화 또는 운영자 검수형 갱신 프로세스 설계"],
            ["Outcome label 구조", "schema/API/dashboard 기록 구현", "실제 사용자 또는 기관 관측 label 수집·import"],
            ["Safety resource", "example YAML", "기관 검토 후 지역/전국 위기 대응 연락처 확정"],
            ["개인정보", "MVP 미구현", "동의, 보존 기간, 삭제 요청, 접근 권한 설계"],
            ["운영자 workflow", "검증 view와 review 입력 구현", "review queue, escalation protocol, audit log 설계"],
            ["Gemini/LLM", "필수 아님", "query expansion, grounded answer generation, operator assistant로 제한적 연결"],
            ["편향 점검", "기초 metric만 있음", "구/군, 접근성, 비용, 대면 부담별 성능 및 안전성 검토"],
        ],
        [2000, 2500, 4860],
    )

    add_heading(doc, "16. Known limitations", 1)
    add_bullets(
        doc,
        [
            "현재 사용자 profile과 학습 label은 synthetic placeholder입니다. feedback/progress schema와 pipeline은 구현되어 있으나, 실제 mission completion, too-hard 피드백, 프로그램 참여, 지원 결과 label은 사용자 또는 기관 관측 후 수집·import해야 합니다.",
            "현재 resource는 공식 출처 기반 seed이지만 자동 크롤링/공공 API 동기화는 아직 없습니다.",
            "지도 거리는 위도/경도 기반 직선거리 근사이며 실제 대중교통 이동시간은 아닙니다.",
            "ML 모델은 임상 위험 모델이 아니며 그렇게 해석하면 안 됩니다.",
            "RAG는 TF-IDF retrieval이므로 생성형 답변 품질이나 semantic matching에는 한계가 있습니다.",
            "실제 배포 전 개인정보, 운영자 권한, 기관 연계, 안전 escalation protocol을 별도로 설계해야 합니다.",
            "추천 정책을 무작위로 흔드는 A/B test는 충분한 안전 데이터와 운영자 검토 체계가 생긴 뒤 제한적으로만 고려해야 합니다.",
        ],
    )

    add_heading(doc, "17. 파일별 검증 상태", 1)
    add_table(
        doc,
        ["검증", "현재 결과"],
        [
            ["make test", "11 passed"],
            ["make eval-sheet", "reports/human_eval_review_sheet.csv 생성 성공"],
            ["make pipeline", "성공. metadata, model card, data card 갱신"],
            ["Dashboard QA", "내 조건→추천 루트→지도→결과 기록 흐름, ID/score 숨김, 모바일/색상 대비 점검"],
        ],
        [2500, 6860],
    )

    add_paragraph(
        doc,
        "Required notice: The current training labels are synthetic placeholders for user behavior and institution outcome labels. The pipeline is implemented, but real observed outcome labels must be collected or imported before production use.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
